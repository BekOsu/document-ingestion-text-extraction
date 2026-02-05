import logging
from pathlib import Path
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: Path) -> dict:
    result = {
        "file_name": file_path.name,
        "page_count": None,
        "extraction_method": "python-docx",
        "extracted_text": None,
        "success": False,
    }

    try:
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        text = "\n".join(paragraphs)
        if text:
            result["extracted_text"] = text
            result["success"] = True
            logger.info(f"Extracted text from {file_path.name}")
        else:
            logger.warning(f"No text found in {file_path.name}")

    except Exception as e:
        logger.error(f"Failed to extract from {file_path.name}: {e}")

    return result